from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SCREENSHOTS = ROOT / "Expected_result_simulation"
OUTPUT = ROOT / "chapter_4_results_and_discussion.docx"

FIGURES = [
    (
        "Figure 4.1: Docker result when the close grey-blue door is detected and the goal-stop threshold is reached",
        SCREENSHOTS / "Screenshot 2026-06-15 100351.png",
    ),
    (
        "Figure 4.2: Docker result when an open grey-blue door region is detected and STOP is selected",
        SCREENSHOTS / "Screenshot 2026-06-15 100531.png",
    ),
    (
        "Figure 4.3: Docker SEARCH result when no valid full door region is detected",
        SCREENSHOTS / "Screenshot 2026-06-15 100606.png",
    ),
    (
        "Figure 4.4: Docker result when a second grey-blue corridor door is detected and STOP is selected",
        SCREENSHOTS / "Screenshot 2026-06-15 100659.png",
    ),
    (
        "Figure 4.5: Docker TURN RIGHT result when the detected door is on the right side of the processed frame",
        SCREENSHOTS / "Screenshot 2026-06-15 101741.png",
    ),
]


def set_cell_shading(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, bold: bool = False, white: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(9)
    if white:
        run.font.color.rgb = RGBColor(255, 255, 255)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(
    document: Document,
    caption: str,
    headers: list[str],
    rows: list[list[str]],
    *,
    page_break_before: bool = False,
) -> None:
    if page_break_before:
        document.add_page_break()
    add_caption(document, caption)
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    header_repeat = OxmlElement("w:tblHeader")
    header_repeat.set(qn("w:val"), "true")
    header_properties.append(header_repeat)
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, "1F4E78")
        set_cell_text(cell, header, bold=True, white=True)
    for row_index, row_data in enumerate(rows):
        row = table.add_row()
        for cell, text in zip(row.cells, row_data):
            set_cell_text(cell, text)
            if row_index % 2 == 0:
                set_cell_shading(cell, "D9EAF7")
    document.add_paragraph()


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, end])


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    add_page_number(section.footer.paragraphs[0])

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for style_name, size in (("Heading 1", 13), ("Heading 2", 12)):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.keep_with_next = True

    caption = document.styles["Caption"]
    caption.font.name = "Times New Roman"
    caption.font.size = Pt(10)
    caption.font.bold = True
    caption.font.italic = False
    caption.font.color.rgb = RGBColor(0, 0, 0)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(6)
    caption.paragraph_format.space_after = Pt(6)
    caption.paragraph_format.keep_with_next = True


def add_title(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run("CHAPTER 4")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(18)
    run = paragraph.add_run("RESULTS AND DISCUSSION")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_body(document: Document, text: str) -> None:
    document.add_paragraph(text)


def add_caption(document: Document, text: str) -> None:
    document.add_paragraph(text, style="Caption")


def add_figure(document: Document, caption: str, path: Path) -> None:
    if not path.exists():
        raise FileNotFoundError(path)
    add_caption(document, caption)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.add_run().add_picture(str(path), width=Inches(6.25))


def build() -> None:
    document = Document()
    configure_document(document)
    add_title(document)

    add_heading(document, "4.1 Introduction")
    add_body(
        document,
        "This chapter presents the observed FYP1 results from the Docker laptop expected-results "
        "demonstration. Unlike the previous expected-results chapter, the discussion in this chapter "
        "is based on screenshots captured while the containerized application was running on the "
        "laptop. The demonstration used the prompt \"find the door\", a browser webcam, and an "
        "OpenCV colour-and-shape detector tuned to the grey-blue corridor doors in the selected test area.",
    )
    add_body(
        document,
        "The Docker demonstration evaluates the high-level software interface before integration with "
        "the Raspberry Pi 4, ESP32, ultrasonic sensors, and motors. Therefore, the screenshots provide "
        "evidence of structured prompt processing, visual door-region detection, and displayed action "
        "selection only. They are not physical robot-navigation results.",
    )

    add_heading(document, "4.2 Docker Demonstration Setup")
    add_body(
        document,
        "The browser captured live webcam frames and sent them to the Docker container. The container "
        "converted the instruction into a structured goal with the target door, the action goal find "
        "and approach the door, low uncertainty, and the grey-blue door colour-and-shape detector. "
        "The processed frame displayed a green bounding box when a valid door region was detected.",
    )
    add_body(
        document,
        "Action selection was based on the detected bounding-box position and apparent image area. "
        "When no valid door region was detected, the application displayed SEARCH. When the detected "
        "door occupied the configured goal-size area, the application displayed STOP. The user moved "
        "the laptop manually; no motor command was executed.",
    )

    add_table(
        document,
        "Table 4.1: Summary of Docker Screenshot Evidence",
        ["Screenshot", "Observed Scene", "Detector Output", "Displayed Action", "Interpretation"],
        [
            [
                "Figure 4.1",
                "Close view of closed grey-blue door 38",
                "Door region enclosed by a green bounding box",
                "STOP",
                "The apparent door area reached the configured goal threshold.",
            ],
            [
                "Figure 4.2",
                "Large open grey-blue door region",
                "Open door region enclosed by a green bounding box",
                "STOP",
                "The detector remained responsive to the door colour and shape under a changed pose.",
            ],
            [
                "Figure 4.3",
                "Corridor view without a valid full grey-blue door region",
                "No bounding box",
                "SEARCH",
                "The application used the safe missing-target fallback.",
            ],
            [
                "Figure 4.4",
                "Second closed grey-blue corridor door",
                "Door region enclosed by a green bounding box",
                "STOP",
                "The detector identified another similarly coloured door in the same test area.",
            ],
            [
                "Figure 4.5",
                "Grey-blue door located on the right side of the frame",
                "Right-side door region enclosed by a green bounding box",
                "TURN RIGHT",
                "The horizontal door position was converted into the corresponding turn action.",
            ],
        ],
        page_break_before=True,
    )

    add_heading(document, "4.3 Observed Results")
    add_heading(document, "4.3.1 Structured Prompt Result", level=2)
    add_body(
        document,
        "Across the screenshots, the application accepted the prompt \"find the door\" and displayed a "
        "structured goal. The visible fields included the original instruction, target door, landmark "
        "door, action goal find and approach the door, low uncertainty, and the selected grey-blue "
        "door colour-and-shape detector. This result demonstrates that the natural-language instruction "
        "was converted into a consistent machine-readable goal before visual processing.",
    )

    add_heading(document, "4.3.2 Door Detection and STOP Result", level=2)
    add_body(
        document,
        "Figure 4.1 shows the closed grey-blue door detected at close range. The green bounding box "
        "covers the main door region and the application displays STOP: expected goal reached. Within "
        "the current rule-based demonstration, this is the expected response because the detected door "
        "occupies a large part of the image.",
    )
    add_figure(document, *FIGURES[0])

    add_body(
        document,
        "Figure 4.2 shows an open grey-blue door region. The detector still forms a large bounding box "
        "and selects STOP. This demonstrates useful tolerance to a changed door pose, but it also shows "
        "that the stopping decision depends on apparent image area rather than measured physical distance.",
    )
    add_figure(document, *FIGURES[1])

    add_heading(document, "4.3.3 Missing-Target SEARCH Result", level=2)
    add_body(
        document,
        "Figure 4.3 shows a corridor frame in which no valid full grey-blue door region satisfies the "
        "detector rules. No green bounding box is produced and the application displays SEARCH: rotate "
        "or move laptop slowly. This is the intended safe fallback because the software does not command "
        "forward movement when the target cannot be grounded.",
    )
    add_figure(document, *FIGURES[2])

    add_heading(document, "4.3.4 Detection of a Second Similar Door", level=2)
    add_body(
        document,
        "Figure 4.4 shows the detector responding to a second grey-blue corridor door. The main door "
        "region is enclosed and STOP is displayed. This screenshot indicates that the current detector "
        "can respond to similarly coloured doors within the same corridor environment, although broader "
        "generalization has not been measured.",
    )
    add_figure(document, *FIGURES[3])

    add_heading(document, "4.3.5 Right-Side Door TURN RIGHT Result", level=2)
    add_body(
        document,
        "Figure 4.5 shows the grey-blue door detected on the right side of the processed frame. "
        "The green bounding box remains on the right of the configured centre zone and the application "
        "displays TURN RIGHT: move laptop right. This provides direct screenshot evidence that the "
        "detector position can be converted into a directional action rather than only SEARCH or STOP.",
    )
    add_figure(document, *FIGURES[4])

    add_heading(document, "4.4 Discussion")
    add_body(
        document,
        "The five screenshots demonstrate the complete Docker software path from a natural-language "
        "instruction to structured goal output, webcam processing, target-region annotation, and a "
        "restricted displayed action. Three screenshot conditions produced STOP after a large grey-blue "
        "door region was detected, one produced SEARCH when no valid region was detected, and one produced "
        "TURN RIGHT when the door was detected on the right side of the processed frame. These counts "
        "describe only the submitted screenshot evidence and must not be interpreted as an accuracy percentage.",
    )
    add_body(
        document,
        "The SEARCH result is important because it demonstrates a conservative response when visual "
        "grounding fails. However, the STOP results also reveal limitations. The system estimates goal "
        "arrival from bounding-box area and does not measure the actual distance to the door. An open "
        "door or nearby similarly coloured structure may therefore trigger STOP. The bounding box may "
        "also include part of the door frame or adjacent structure.",
    )
    add_body(
        document,
        "SEARCH, TURN RIGHT, and STOP are directly supported by the submitted screenshots. TURN LEFT and "
        "MOVE FORWARD remain implemented actions, but they require separate recorded evidence before their "
        "performance can be discussed as observed results. The detector is tuned to grey-blue doors in one "
        "corridor and is not a general semantic door-recognition model.",
    )

    add_table(
        document,
        "Table 4.2: Interpretation and Limitations of the Demonstrated Results",
        ["Area", "Evidence from Screenshots", "Limitation"],
        [
            [
                "Prompt processing",
                "The instruction is represented as a structured door goal with low uncertainty.",
                "Only one supported instruction is demonstrated.",
            ],
            [
                "Visual grounding",
                "Green bounding boxes identify grey-blue door regions in three screenshots.",
                "Detector is tuned to colour and geometry in the selected corridor.",
            ],
            [
                "Safe fallback",
                "SEARCH is displayed when no valid door region is detected.",
                "Physical search rotation is not demonstrated.",
            ],
            [
                "Directional action",
                "TURN RIGHT is displayed when the detected door is on the right side of the frame.",
                "TURN LEFT and physical turning are not yet demonstrated.",
            ],
            [
                "Goal stopping",
                "STOP is displayed when the detected region occupies a large image area.",
                "Apparent area is not a direct physical-distance measurement.",
            ],
            [
                "Robot execution",
                "No physical movement is claimed.",
                "Raspberry Pi, ESP32, sensors, motors, and safety overrides remain to be validated.",
            ],
        ],
    )

    add_heading(document, "4.5 Required Physical Validation During FYP2")
    add_body(
        document,
        "The Docker results establish a working software demonstration, but the final project result must "
        "come from the Raspberry Pi 4 physical robot. The same prompt and detector should be transferred "
        "to the Raspberry Pi and connected to the ESP32 safety and motor-control layer. Repeated trials "
        "are required before reporting accuracy, latency, movement success, or stopping performance.",
    )
    add_table(
        document,
        "Table 4.3: Required FYP2 Validation After the Docker Demonstration",
        ["Evaluation Area", "Required Measurement"],
        [
            ["Door detection", "Repeated detection and missing-target trials under different lighting, distance, and door poses"],
            ["Action selection", "Recorded SEARCH, TURN LEFT, TURN RIGHT, MOVE FORWARD, and STOP correctness"],
            ["Raspberry Pi performance", "Frame-processing rate and end-to-end action latency"],
            ["ESP32 communication", "USB serial command and sensor-status reliability"],
            ["Safety", "Ultrasonic obstacle stop, sensor-fault stop, invalid-command stop, and timeout stop"],
            ["Physical navigation", "Door-approach success rate and measured stopping distance"],
            ["Power system", "Voltage stability, resets, current, and operating duration"],
        ],
        page_break_before=True,
    )

    add_heading(document, "4.6 Summary")
    add_body(
        document,
        "This chapter presented the observed FYP1 Docker demonstration results for the prompt \"find the "
        "door\". The screenshots show consistent structured goal output, grey-blue door-region detection, "
        "STOP selection for large detected door regions, SEARCH selection when the target is not grounded, "
        "and TURN RIGHT selection when the door is on the right side of the frame. The demonstration "
        "confirms that the high-level software interface is functional, while also showing that physical-"
        "distance estimation, broader door recognition, remaining action evidence, and real robot "
        "validation remain necessary during FYP2.",
    )

    document.save(OUTPUT)
    print(f"Built {OUTPUT}")


if __name__ == "__main__":
    build()
