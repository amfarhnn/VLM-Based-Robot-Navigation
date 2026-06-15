from __future__ import annotations

import re
import shutil
import tempfile
import zipfile
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
TARGET = ROOT / "i_editted.docx"
SOURCE = ROOT / "thesis_chapters_1_to_4.docx"
BACKUP = ROOT / "i_editted_before_raspberry_update.docx"


def replace_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def has_image(paragraph) -> bool:
    return bool(paragraph._element.xpath(".//a:blip"))


def is_caption(text: str) -> bool:
    return bool(re.match(r"^(Figure|Table)\s+\d+\.\d+:", text.strip()))


def is_numbered_heading(text: str) -> bool:
    return bool(re.match(r"^\d+\.\d+(?:\.\d+)?(?:\s|$)", text.strip()))


def find_heading(document: Document, number: str):
    pattern = re.compile(rf"^{re.escape(number)}(?:\s|$)")
    for paragraph in document.paragraphs:
        if pattern.match(paragraph.text.strip()):
            return paragraph
    raise ValueError(f"Heading {number} was not found")


def paragraph_index(paragraphs, paragraph) -> int:
    for index, candidate in enumerate(paragraphs):
        if candidate._p is paragraph._p:
            return index
    raise ValueError("Paragraph is not present in the document paragraph list")


def section_body(document: Document, number: str):
    paragraphs = document.paragraphs
    start = paragraph_index(paragraphs, find_heading(document, number))
    body = []
    for paragraph in paragraphs[start + 1 :]:
        text = paragraph.text.strip()
        if text.startswith("CHAPTER ") or is_numbered_heading(text) or text == "REFERENCE":
            break
        if not text or is_caption(text) or has_image(paragraph):
            continue
        body.append(paragraph)
    return body


def replace_section_body(target: Document, source: Document, number: str) -> None:
    target_body = section_body(target, number)
    source_text = [paragraph.text for paragraph in section_body(source, number)]
    for paragraph, text in zip(target_body, source_text):
        replace_paragraph_text(paragraph, text)
    for paragraph in target_body[len(source_text) :]:
        delete_paragraph(paragraph)
    if len(source_text) > len(target_body):
        anchor = find_heading(target, number)
        paragraphs = target.paragraphs
        start = paragraph_index(paragraphs, anchor)
        next_heading = None
        for paragraph in paragraphs[start + 1 :]:
            text = paragraph.text.strip()
            if text.startswith("CHAPTER ") or is_numbered_heading(text) or text == "REFERENCE":
                next_heading = paragraph
                break
        if next_heading is None:
            raise ValueError(f"Could not find insertion point after {number}")
        template = target_body[-1] if target_body else anchor
        for text in source_text[len(target_body) :]:
            new_paragraph = deepcopy(template._p)
            next_heading._p.addprevious(new_paragraph)
            refreshed = target.paragraphs
            replace_paragraph_text(
                refreshed[paragraph_index(refreshed, next_heading) - 1],
                text,
            )


def replace_exact(document: Document, old: str, new: str) -> None:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == old:
            replace_paragraph_text(paragraph, new)
            return
    raise ValueError(f"Paragraph not found: {old}")


def delete_exact(document: Document, text: str) -> None:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text:
            delete_paragraph(paragraph)
            return
    raise ValueError(f"Paragraph not found for deletion: {text}")


def replace_text_everywhere(document: Document, old: str, new: str) -> int:
    count = 0
    for paragraph in document.paragraphs:
        if old in paragraph.text:
            replace_paragraph_text(paragraph, paragraph.text.replace(old, new))
            count += 1
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old in paragraph.text:
                        replace_paragraph_text(paragraph, paragraph.text.replace(old, new))
                        count += 1
    return count


def image_part_name(paragraph) -> str | None:
    blips = paragraph._element.xpath(".//a:blip")
    if not blips:
        return None
    relation_id = blips[0].get(
        "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
    )
    if not relation_id:
        return None
    return Path(paragraph.part.rels[relation_id].target_part.partname).name


def insert_caption_before_image(
    document: Document, image_name: str, caption: str, caption_template
) -> None:
    if any(paragraph.text.strip() == caption for paragraph in document.paragraphs):
        return
    for paragraph in document.paragraphs:
        if image_part_name(paragraph) == image_name:
            clone = deepcopy(caption_template._p)
            paragraph._p.addprevious(clone)
            replace_paragraph_text(Paragraph(clone, paragraph._parent), caption)
            return
    raise ValueError(f"Image {image_name} was not found")


def parse_markdown_tables(path: Path) -> dict[str, list[list[str]]]:
    lines = path.read_text(encoding="utf-8").splitlines()
    tables: dict[str, list[list[str]]] = {}
    for index, line in enumerate(lines):
        match = re.fullmatch(r"\*\*(Table \d+\.\d+:[^*]+)\*\*", line.strip())
        if not match:
            continue
        caption = match.group(1).strip()
        rows: list[list[str]] = []
        cursor = index + 1
        while cursor < len(lines) and not lines[cursor].strip():
            cursor += 1
        while cursor < len(lines) and lines[cursor].strip().startswith("|"):
            cells = [cell.strip().replace("`", "") for cell in lines[cursor].strip().strip("|").split("|")]
            if not all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
                rows.append(cells)
            cursor += 1
        tables[caption] = rows
    return tables


def set_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    replace_paragraph_text(paragraph, text)
    for extra in list(cell.paragraphs[1:]):
        delete_paragraph(extra)


def set_table(table, rows: list[list[str]]) -> None:
    while len(table.rows) < len(rows):
        table._tbl.append(deepcopy(table.rows[-1]._tr))
    while len(table.rows) > len(rows):
        table._tbl.remove(table.rows[-1]._tr)
    expected_columns = len(table.columns)
    for row_data in rows:
        if len(row_data) != expected_columns:
            raise ValueError(
                f"Table expects {expected_columns} columns but received {len(row_data)}: {row_data}"
            )
    for row, row_data in zip(table.rows, rows):
        for cell, text in zip(row.cells, row_data):
            set_cell_text(cell, text)


def system_image_replacements() -> dict[str, Path]:
    return {
        "image12.png": ROOT / "drawio/chapter_2/exported/figure_2_1_general_language_guided_navigation_process.png",
        "image2.png": ROOT / "drawio/chapter_2/exported/figure_2_2_lm_nav_baseline_pipeline.png",
        "image1.png": ROOT / "drawio/chapter_2/exported/figure_2_3_positioning_of_related_works_around_lm_nav.png",
        "image4.png": ROOT / "drawio/chapter_3/exported/figure_3_1_proposed_project_workflow.png",
        "image8.png": ROOT / "drawio/chapter_3/exported/figure_3_2_proposed_system_architecture.png",
        "image17.png": ROOT / "figures/chapter_3/figure_3_3_main_hardware_components.png",
        "image6.png": ROOT / "figures/chapter_3/figure_3_4_finalized_fusion_360_robot_model.png",
        "image16.png": ROOT / "figures/chapter_3/figure_3_5_fusion_360_robot_multi_view.png",
        "image7.png": ROOT / "figures/chapter_3/figure_3_6_mechanical_design_sketch_and_dimensions.png",
        "image19.png": ROOT / "drawio/chapter_3/exported/figure_3_7_finalized_component_based_circuit_architecture.png",
        "image11.png": ROOT / "drawio/chapter_3/exported/figure_3_8_finalized_esp32_sensor_and_motor_wiring.png",
        "image9.png": ROOT / "figures/chapter_4/figure_4_1_expected_obstacle_stop_response.png",
        "image5.png": ROOT / "figures/chapter_4/figure_4_2_expected_command_timeout_response.png",
    }


def replace_docx_media(source_docx: Path, output_docx: Path) -> None:
    replacements = system_image_replacements()
    for path in replacements.values():
        if not path.exists():
            raise FileNotFoundError(path)
    with tempfile.TemporaryDirectory() as temporary:
        temporary_path = Path(temporary)
        extracted = temporary_path / "docx"
        with zipfile.ZipFile(source_docx) as archive:
            archive.extractall(extracted)
        media = extracted / "word" / "media"
        for name, replacement in replacements.items():
            shutil.copyfile(replacement, media / name)
        with zipfile.ZipFile(output_docx, "w", zipfile.ZIP_DEFLATED) as archive:
            for path in extracted.rglob("*"):
                if path.is_file():
                    archive.write(path, path.relative_to(extracted).as_posix())


def custom_tables() -> dict[int, list[list[str]]]:
    return {
        7: [
            ["Component", "Finalized Role", "Main Interface", "Main Strength", "Main Constraint"],
            ["Raspberry Pi 4", "Onboard prompt processing, webcam capture, OpenCV grounding, action selection, and logging", "USB webcam and USB serial", "Flexible Linux and Python environment", "Limited performance for large models"],
            ["ESP32", "Sensor reading, deterministic safety, and four-motor control", "USB serial and GPIO", "Independent real-time safety layer", "Runs only approved low-level commands"],
            ["USB webcam", "Front RGB image input", "Raspberry Pi USB", "Low-cost visual input", "No dense depth or 3D map"],
            ["Two HC-SR04 sensors", "Front-left and front-right obstacle safety", "ESP32 GPIO with Echo level shifting", "Low-cost local obstacle detection", "Limited coverage and accuracy"],
            ["GY-291 / ADXL345", "Acceleration, roll/pitch tilt, motion, vibration, and shock observations", "ESP32 I2C", "Useful chassis-condition feedback", "Does not provide yaw heading"],
            ["Two MX1508 drivers", "Four independent DC motor channels", "ESP32 GPIO and regulated motor rail", "Compact four-motor control", "Requires current and thermal verification"],
            ["Four DC gear motors", "Physical forward, turn, search, and stop behaviour", "MX1508 outputs", "Simple four-wheel movement", "No wheel-encoder feedback"],
        ],
        9: [
            ["Scenario", "Example Instruction", "Expected Behaviour", "Purpose"],
            ["Supported target absent", '"Find the green marker."', "search", "Verify safe visual search"],
            ["Target on left", '"Find the green marker."', "turn_left", "Verify left action selection"],
            ["Target on right", '"Find the green marker."', "turn_right", "Verify right action selection"],
            ["Target centred and distant", '"Find the green marker."', "move_forward", "Verify approach action"],
            ["Target centred and close", '"Find the green marker."', "stop", "Verify goal-stop behaviour"],
            ["Unsupported or ambiguous target", '"Go there."', "stop", "Verify prompt uncertainty handling"],
            ["Obstacle present", '"Find the green marker."', "ESP32 forces stop", "Verify independent local safety"],
        ],
        12: [
            ["Prompt Input", "First-Prototype Interpretation", "Expected Result"],
            ['"Find the green marker."', "Supported target and basic goal", "Structured low-uncertainty output"],
            ['"Go to the green marker."', "Supported target and approach goal", "Structured low-uncertainty output"],
            ['"Go there."', "Unsupported ambiguous target", "High uncertainty and stop"],
            ["Unsupported landmark instruction", "Target outside first-prototype capability", "High uncertainty and stop"],
        ],
        14: [
            ["Condition", "Selected Action", "Reason"],
            ["Structured target unsupported or uncertain", "stop", "Prevent unsupported execution"],
            ["Marker not detected", "search", "Rotate slowly to locate target"],
            ["Marker centre is left of image centre zone", "turn_left", "Align camera with target"],
            ["Marker centre is right of image centre zone", "turn_right", "Align camera with target"],
            ["Marker is centred and below goal-size threshold", "move_forward", "Approach the target"],
            ["Marker reaches configured image-area threshold", "stop", "Target is considered reached"],
            ["ESP32 reports obstacle or sensor fault", "stop", "Local safety overrides high-level action"],
        ],
        16: [
            ["Criterion", "Raspberry Pi 4 Baseline", "Future Extension"],
            ["Physical implementation", "One Raspberry Pi 4 and ESP32 robot", "Retain the same safe robot base"],
            ["Prompt processing", "Deterministic fixed-schema parser", "Lightweight local model or API-based VLM"],
            ["Visual grounding", "OpenCV HSV coloured-marker detection", "Object detector or vision-language grounding"],
            ["Navigation action", "Restricted five-action set", "Relation-aware planning or learned policy"],
            ["Position feedback", "Ultrasonic and GY-291 observations", "Wheel encoders, magnetometer, or full IMU"],
            ["Environment model", "Current RGB frame only", "RGB-D, LiDAR, mapping, or scene graph"],
            ["Result role", "Final physical robot results", "More advanced FYP2 or future research"],
        ],
        17: [
            ["Component or Tool", "Company or Organization", "Involvement"],
            ["Raspberry Pi 4", "Raspberry Pi Ltd", "Only onboard high-level compute board"],
            ["ESP32", "Espressif Systems", "Low-level sensing, safety, and motor control"],
            ["GY-291 / ADXL345", "Analog Devices device on third-party module", "Acceleration, roll/pitch tilt, motion, vibration, and shock sensing"],
            ["MX1508 modules", "Supplier to be confirmed", "Four-motor driver subsystem"],
            ["Autodesk Fusion 360", "Autodesk", "Mechanical design"],
            ["Python", "Python Software Foundation", "High-level software"],
            ["OpenCV", "OpenCV project", "Webcam processing and marker grounding"],
            ["Arduino framework", "Arduino project / Espressif ecosystem", "ESP32 firmware"],
            ["Docker", "Docker, Inc.", "Isolated laptop expected-results demonstration"],
            ["diagrams.net / draw.io", "JGraph / diagrams.net", "Report diagrams"],
        ],
        24: [
            ["Area", "Raspberry Pi 4 Physical Prototype", "Laptop Docker Expected-Results Demonstration"],
            ["Camera input", "Robot-mounted USB webcam", "Laptop browser webcam"],
            ["Processing", "Onboard Python and OpenCV", "Containerized Flask, Python, OpenCV, and NumPy"],
            ["Action output", "Validated command sent to ESP32", "Instruction displayed on screen for manual movement"],
            ["Result status", "Final physical robot result during FYP2", "Expected-results evidence only"],
        ],
    }


def main() -> None:
    if not TARGET.exists() or not SOURCE.exists():
        raise FileNotFoundError("Both i_editted.docx and thesis_chapters_1_to_4.docx are required")
    if not BACKUP.exists():
        shutil.copy2(TARGET, BACKUP)

    target = Document(TARGET)
    source = Document(SOURCE)

    for chapter_section in [
        "1.1", "1.2", "1.3", "1.4", "1.5", "1.6", "1.7", "1.8", "1.9",
        "2.1", "2.2", "2.3", "2.4", "2.5", "2.5.1", "2.5.2", "2.5.3",
        "2.5.4", "2.6", "2.7", "2.8", "2.9", "2.10", "2.11", "2.12", "2.13",
        "4.1", "4.2", "4.3", "4.4", "4.5", "4.6", "4.7",
    ]:
        replace_section_body(target, source, chapter_section)

    exact_replacements = {
        "3.1 Introduction": "3.1 Introduction",
        "This chapter presents the finalized methodology for the project Prompt Engineering for Mobile Robot Navigation. The project investigates how structured prompt engineering and webcam-based visual grounding can be used to select safe, simple movement actions for a low-cost indoor mobile robot.":
            "This chapter presents the finalized methodology for Prompt Engineering for Mobile Robot Navigation. The project contains one physical implementation: a Raspberry Pi 4 mobile robot with an ESP32 sensor, safety, and motor-control layer.",
        "The methodology is inspired by the modular structure of LM-Nav and the prompt-based action-selection idea of VLMnav. LM-Nav motivates the separation between language understanding, visual grounding, and navigation execution. VLMnav motivates restricting model output to a fixed action set.":
            "The methodology is inspired by the modular structure of LM-Nav and the prompt-based action-selection idea of VLMnav. Language interpretation, visual grounding, and navigation execution are separated so each stage can be tested. The first prototype targets find the green marker and restricts all output to the approved action set.",
        "The finalized project contains exactly two physical robot approaches:":
            "The finalized project contains one Raspberry Pi 4 physical robot approach.",
        "- **Approach 1:** Raspberry Pi 4 and ESP32.": "- Raspberry Pi 4 performs onboard prompt processing, webcam processing, action selection, and logging.",
        "- **Approach 2:** Coral Dev Board and ESP32.": "- ESP32 performs sensor reading, deterministic safety, and four-motor control.",
        "Both approaches use the same four-wheel chassis, USB webcam, two MX1508 motor drivers, adjustable motor buck converter, two HC-SR04 ultrasonic sensors, GY-291 / ADXL345 accelerometer, three 18650 cells in series, and a 3S-compatible power-bank charging/output module. Only the high-level compute board and compatible model runtime change.":
            "The robot uses a four-wheel chassis, USB webcam, two MX1508 motor drivers, adjustable motor buck converter, two HC-SR04 ultrasonic sensors, GY-291 / ADXL345 accelerometer, and the completed protected power system.",
        "The main contribution is the design, implementation, and comparison of two onboard prompt-engineered indoor navigation systems that share the same low-level robot platform.":
            "The laptop-only Docker demonstration produces expected-results evidence before physical testing. It is not a second physical approach.",
        "The project follows an iterative prototyping and comparative experimental methodology. Shared hardware modules are tested first, followed by the Raspberry Pi 4 implementation and Coral Dev Board implementation. The two approaches are then compared using the same instructions, indoor landmarks, action set, safety rules, and evaluation metrics.":
            "The project follows an iterative prototype methodology. The already-built mechanical and power systems are completed by wiring and validating the ESP32 sensor and motor layer, connecting Raspberry Pi USB serial, verifying webcam capture and OpenCV visual grounding, and testing one complete basic goal.",
        "3.3 Proposed System Architecture": "3.3 Finalized System Architecture",
        "Both approaches use the same layered architecture. The high-level compute board processes language and webcam input. The ESP32 handles real-time sensors, safety, and motor control. This division prevents model latency from directly controlling motor timing.":
            "The Raspberry Pi 4 performs high-level processing. The ESP32 handles local real-time sensing, safety, and motor control. USB serial is used between the boards because it is simple to debug and leaves all ESP32 motor GPIO pins available.",
        "3.3.1 Finalized Implementation Approaches and Hardware Roles": "3.3.1 Hardware Roles",
        "Table 3.3: Finalized Implementation Approaches and Hardware Roles": "Table 3.3: Finalized Hardware Roles",
        "3.4 System Requirements, Constraints, and Acceptance Criteria": "3.4 Requirements, Constraints, and Acceptance Criteria",
        "The baseline is limited to controlled indoor areas and a small action set. It does not include full SLAM, LiDAR mapping, dense depth, or a learned navigation policy.":
            "The baseline does not include SLAM, LiDAR mapping, dense depth, or an unrestricted large vision-language model.",
        "Table 3.5: Instruction Categories and Indoor Landmarks": "Table 3.5: First-Prototype Scenarios",
        "Figure 3.3: Main hardware components for the finalized approaches": "Figure 3.3: Main hardware components for the finalized Raspberry Pi 4 robot",
        "Prompt engineering converts user instructions into a structured navigation representation. Both approaches use the same schema so that compute-platform performance can be compared fairly.":
            "The first prototype uses a deterministic structured parser. This applies the main prompt-engineering principles required by the study: fixed fields, restricted actions, explicit uncertainty, and safe rejection of unsupported goals.",
        "Table 3.8: Prompt Templates for Evaluation": "Table 3.8: Supported First-Prototype Prompt Inputs",
        "Table 3.9: Structured Output Fields": "Table 3.9: Structured Prompt Fields",
        "After prompt processing, the selected compute board captures a USB-webcam frame and combines it with the structured instruction. The Raspberry Pi 4 uses a lightweight onboard method. The Coral Dev Board uses a compatible quantized TFLite or Edge TPU model.":
            "The Raspberry Pi captures a USB-webcam frame and detects the selected marker using OpenCV HSV colour segmentation. The largest valid marker region is used for simple action selection.",
        "Table 3.10: Visual Grounding and Action Selection Strategies": "Table 3.10: Visual Grounding and Action Decision Rules",
        "3.8 Robot Action Set and Motor Command Mapping": "3.8 Robot Action and Command Mapping",
        "Table 3.11: Action Set and Motor Command Mapping": "Table 3.11: Action Set and Motor Command Mapping",
        "The GY-291 is used to measure X/Y/Z acceleration, calculate roll and pitch tilt, and observe motion, vibration, and shock during robot operation.":
            "The ESP32 accepts only the approved commands. Unknown commands, loss of both ultrasonic readings, and a command delay above 1,000 ms cause stop.",
        "3.9 Approach Comparison Criteria": "3.9 Baseline and Future Extension Boundaries",
        "Table 3.12: Comparison of Implementation Approaches": "Table 3.12: Baseline and Future Extension Boundaries",
        "The finalized Fusion 360 robot design provides mounting positions for the selected compute board, ESP32, USB webcam, two front ultrasonic sensors, GY-291, two MX1508 drivers, battery system, power modules, and four motors. The same chassis is used for both approaches so that the compute-board comparison is not affected by different mechanical layouts.":
            "The finalized Fusion 360 robot design provides mounting positions for the Raspberry Pi 4, ESP32, USB webcam, two front ultrasonic sensors, GY-291, two MX1508 drivers, completed power system, and four motors.",
        "The overview render shows the front-mounted webcam and ultrasonic sensors, central protected battery position, four-wheel drive arrangement, and elevated compute-board mounting area. The supplied render illustrates the Raspberry Pi 4 configuration; Approach 2 replaces it with the Coral Dev Board while retaining the same shared chassis and compute-board mounting area.":
            "The existing render records the physical design and Raspberry Pi 4 mounting layout.",
        "The front, side, rear, and top views are used to verify sensor visibility, wheel clearance, component spacing, cable access, and the mounting locations for both compute-board approaches.":
            "The front, side, rear, and top views support sensor visibility, wheel clearance, component spacing, cable access, and Raspberry Pi 4 mounting verification.",
        "The robot design includes modular printable parts for the Raspberry Pi 4, ESP32, battery system, motor drivers, GY-291, webcam, and both ultrasonic sensors. Separate left and right ultrasonic mounts position the sensors at the front of the chassis, while the webcam fitting and stand provide a fixed forward-facing camera location. The compute-board, IMU, battery, and driver mounts organize components and reduce movement during operation.":
            "The robot design includes modular printable parts for the Raspberry Pi 4, ESP32, battery system, motor drivers, GY-291, webcam, and both ultrasonic sensors. These mounts organize components, maintain sensor visibility, and reduce movement during operation.",
        "The STL files are stored in the project STL/ directory so that damaged or revised mounts can be reproduced without redesigning the complete chassis. Approach 2 may require an adjusted compute-board mounting plate for the exact Coral Dev Board dimensions while retaining the same shared chassis locations.":
            "The STL files are stored in the project STL/ directory so that damaged or revised mounts can be reproduced without redesigning the complete chassis.",
        "3.11 Company, Software, and Tool Involvement": "3.11 Company, Software, and Tool Involvement",
        "Table 3.14: Source Code Modules for Baseline Implementation": "Table 3.14: Source Code and Documentation Modules",
        "3.17 Safety, Ethics, and Practical Considerations": "3.17 Safety, Limitations, and Future Work",
        "Testing must use matched protected cells, an appropriate 3S battery-management system, a fuse, and a reachable main switch. The selected power-bank module must explicitly support 3S input. Motor tests must begin with the robot raised from the floor. HC-SR04 Echo pins require level shifting for ESP32 safety.":
            "Testing must use the completed protected power system, fuse, reachable main switch, correct regulated rails, and a common ground. Motor tests begin with the robot raised. HC-SR04 Echo pins require level shifting. During the first prototype, the ESP32 is powered only by the Raspberry Pi USB cable.",
        "The robot must stop for invalid output, high uncertainty, obstacle detection, unknown command, or command timeout. Tests must be supervised in a controlled indoor area. Webcam testing should avoid recording identifiable people without permission.":
            "The robot stops for invalid output, high uncertainty, obstacle detection, sensor fault, unknown commands, or command timeout. Testing is supervised in a controlled indoor area, and webcam testing avoids recording identifiable people without permission.",
        "The webcam does not provide dense depth or a 3D map. Two front ultrasonic sensors provide limited obstacle coverage. GY-291 roll and pitch calculations are most reliable while the robot is stationary or moving slowly because dynamic acceleration and vibration affect gravity-based tilt calculations. The Raspberry Pi 4 cannot run large models efficiently, while the Coral Dev Board supports only compatible Edge TPU models.":
            "The USB webcam does not provide dense depth or a 3D map. Two ultrasonic sensors provide limited obstacle coverage. The GY-291 does not provide yaw heading. The Raspberry Pi 4 has limited performance for large models, and the first prototype supports only simple coloured-marker goals.",
        "Future improvements may include wheel encoders, RGB-D camera, LiDAR, stronger motor drivers, improved battery monitoring, mapping, ViNT, NoMaD, or more advanced VLM/VLA models.":
            "Future work may include lightweight object detection, API-based VLM experiments, wheel encoders, magnetometer or full IMU, RGB-D camera, LiDAR, mapping, relation-aware prompts, and stronger navigation policies.",
        "3.19 Summary": "3.18 Summary",
        "This chapter finalized the project around two physical onboard approaches: Raspberry Pi 4 with ESP32 and Coral Dev Board with ESP32. Both approaches share the same four-wheel chassis, two MX1508 motor drivers, adjustable motor buck converter, two front HC-SR04 sensors, GY-291 accelerometer, USB webcam, 3S battery pack, and 3S-compatible power module.":
            "This chapter finalized the methodology around one Raspberry Pi 4 physical robot. The project connects structured prompt processing and OpenCV visual grounding to an ESP32 that independently handles sensors, safety, and four-motor control.",
        "The methodology defines a fair platform comparison using common prompts, scenarios, safety rules, hardware, and evaluation metrics. It also separates high-level language and visual processing from ESP32 real-time sensing and motor control.":
            "The first measurable target is a reliable and safe single-goal indoor marker-navigation demonstration.",
    }
    for old, new in exact_replacements.items():
        if old != new:
            replace_exact(target, old, new)
    delete_exact(target, "3.18 Limitations and Future Work")

    caption_updates = {
        "Figure 2.2: LM-Nav baseline pipeline": "Figure 2.3: LM-Nav baseline pipeline",
        "Figure 2.3: Positioning of related works around LM-Nav": "Figure 2.12: Positioning of related works around LM-Nav",
        "Table 3.7: Software, Model, and Platform Requirements": "Table 3.7: Software and Model Requirements",
        "Table 3.13: Company, Component, and Software Involvement": "Table 3.13: Company, Component, and Software Involvement",
        "Table 3.15: Main Equations and Decision Rules": "Table 3.15: Main Equations and Decision Rules",
        "Table 3.16: Implementation Procedure": "Table 3.16: Implementation Procedure",
        "Table 3.17: Testing and Validation Matrix": "Table 3.17: Testing and Validation Matrix",
        "Table 3.18: Evaluation Metrics": "Table 3.18: Evaluation Metrics",
        "4.2 Expected System Results": "4.2 Expected Basic Prototype Result",
        "Table 4.1: Summary of Expected System Results": "Table 4.1: Summary of Expected Basic Prototype Results",
        "4.3 Expected Platform Comparison": "4.3 Expected Docker Laptop Demonstration",
        "Table 4.2: Expected Raspberry Pi 4 and Coral Dev Board Comparison": "Table 4.2: Expected Raspberry Pi Physical and Docker Demonstration Roles",
        "Table 4.3: Required FYP2 Measurements": "Table 4.3: Required FYP2 Measurements",
    }
    for old, new in caption_updates.items():
        if old != new:
            replace_exact(target, old, new)

    replace_text_everywhere(target, "This pipeline is shown in Figure 2.2.", "This pipeline is shown in Figure 2.3.")
    replace_text_everywhere(
        target,
        "The positioning of the reviewed papers around the proposed baseline is shown in Figure 2.3.",
        "The positioning of the reviewed papers around the proposed baseline is shown in Figure 2.12.",
    )

    caption_template = next(
        paragraph for paragraph in target.paragraphs
        if paragraph.text.strip() == "Figure 2.1: General language-guided navigation process"
    )
    paper_captions = {
        "image20.png": "Figure 2.2: LM-Nav real-world visual navigation platform (Shah et al., 2023a)",
        "image15.png": "Figure 2.4: VLMnav prompt and action-selection interface (Goetting et al., 2024)",
        "image18.png": "Figure 2.5: VLMaps language-grounded navigation example (Huang et al., 2023)",
        "image22.png": "Figure 2.6: HOV-SG hierarchical navigation platform (Werby et al., 2024)",
        "image13.png": "Figure 2.7: ViNT visual navigation transformer architecture (Shah et al., 2023b)",
        "image21.png": "Figure 2.8: NoMaD visual navigation examples (Sridhar et al., 2023)",
        "image14.png": "Figure 2.9: NaVid video-based navigation architecture (Zhang et al., 2024)",
        "image23.png": "Figure 2.10: Uni-NaVid embodied navigation examples (Zhang et al., 2026)",
        "image10.png": "Figure 2.11: NaVILA vision-language-action navigation examples (Cheng et al., 2025)",
    }
    for image_name, caption in paper_captions.items():
        insert_caption_before_image(target, image_name, caption, caption_template)

    markdown_tables = {}
    markdown_tables.update(parse_markdown_tables(ROOT / "literature_review.md"))
    markdown_tables.update(parse_markdown_tables(ROOT / "chapter_3_methodology.md"))
    markdown_tables.update(parse_markdown_tables(ROOT / "chapter_4_results_and_discussion.md"))
    table_caption_map = {
        1: "Table 2.1: Roles of prompt engineering in mobile robot navigation",
        2: "Table 2.2: Critical comparison of related papers",
        3: "Table 2.3: Hardware and method suitability for this FYP",
        4: "Table 2.4: Research gap synthesis and project strategy",
        5: "Table 3.1: Methodology Alignment with Research Objectives",
        6: "Table 3.2: Development Workflow",
        8: "Table 3.4: Requirements and Acceptance Criteria",
        10: "Table 3.6: Experimental Materials",
        11: "Table 3.7: Software, Model, and Platform Requirements",
        13: "Table 3.8: Structured Prompt Fields",
        15: "Table 3.10: Action Set and Motor Command Mapping",
        18: "Table 3.11: Source Code and Documentation Modules",
        19: "Table 3.12: Main Equations and Decision Rules",
        20: "Table 3.13: Implementation Procedure",
        21: "Table 3.14: Testing and Validation Matrix",
        22: "Table 3.15: Evaluation Metrics",
        23: "Table 4.1: Summary of Expected Basic Prototype Results",
        25: "Table 4.2: Required FYP2 Measurements",
    }
    custom = custom_tables()
    for index, table in enumerate(target.tables, start=1):
        if index in custom:
            rows = custom[index]
        else:
            caption = table_caption_map[index]
            rows = markdown_tables[caption]
        set_table(table, rows)

    with tempfile.TemporaryDirectory() as temporary:
        intermediate = Path(temporary) / "updated.docx"
        final = Path(temporary) / "updated_with_images.docx"
        target.save(intermediate)
        replace_docx_media(intermediate, final)
        shutil.copy2(final, TARGET)

    print(f"Updated {TARGET.name}")
    print(f"Backup: {BACKUP.name}")


if __name__ == "__main__":
    main()
